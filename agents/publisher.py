"""Publisher Agent - Converts reports to various output formats."""

from typing import Dict, Any
from agents.base_agent import BaseAgent
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from pathlib import Path
import config


class PublisherAgent(BaseAgent):
    """
    Exports research reports to PDF, DOCX, and Markdown formats.
    """
    
    def __init__(self, llm_client=None):
        super().__init__("PublisherAgent", llm_client)
        
    async def execute(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Export report to requested format(s).
        
        Args:
            input_data: {
                "report_with_citations": str,
                "metadata": Dict,
                "output_formats": List[str] - ["pdf", "docx", "markdown"],
                "output_filename": str (optional)
            }
            
        Returns:
            {
                "output_files": Dict[str, str] - format: filepath,
                "success": bool
            }
        """
        self.validate_input(input_data, ["report_with_citations"])
        
        report = input_data["report_with_citations"]
        metadata = input_data.get("metadata", {})
        formats = input_data.get("output_formats", ["pdf", "markdown"])
        filename = input_data.get("output_filename", "research_report")
        
        # Sanitize filename
        filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_'))
        filename = filename.replace(' ', '_')
        
        self.log_progress(f"Publishing to formats: {', '.join(formats)}")
        
        output_files = {}
        
        # Generate each format
        if "markdown" in formats or "md" in formats:
            md_path = await self._export_markdown(report, filename)
            output_files["markdown"] = str(md_path)
            
        if "pdf" in formats:
            pdf_path = await self._export_pdf(report, metadata, filename)
            output_files["pdf"] = str(pdf_path)
            
        if "docx" in formats:
            docx_path = await self._export_docx(report, metadata, filename)
            output_files["docx"] = str(docx_path)
            
        self.log_progress(f"Published {len(output_files)} formats")
        
        return {
            "output_files": output_files,
            "success": len(output_files) > 0
        }
        
    async def _export_markdown(self, report: str, filename: str) -> Path:
        """Export to Markdown file."""
        
        output_path = config.REPORTS_DIR / f"{filename}.md"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report)
            
        self.log_progress(f"Saved Markdown: {output_path}")
        return output_path
        
    async def _export_pdf(
        self,
        report: str,
        metadata: Dict,
        filename: str
    ) -> Path:
        """Export to PDF using ReportLab."""
        
        output_path = config.REPORTS_DIR / f"{filename}.pdf"
        
        # Create PDF
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )
        
        # Styles
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#2C3E50',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor='#34495E',
            spaceAfter=12,
            spaceBefore=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        )
        
        # Build content
        story = []
        
        # Title page
        title = metadata.get("title", "Research Report")
        story.append(Spacer(1, 2*inch))
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(
            f"Generated: {metadata.get('date', 'N/A')}",
            styles['Normal']
        ))
        story.append(PageBreak())
        
        # Parse report content
        for line in report.split('\n'):
            line = line.strip()
            
            if not line:
                story.append(Spacer(1, 0.2*inch))
                continue
                
            # Handle markdown headings
            if line.startswith('# '):
                text = line.lstrip('#').strip()
                story.append(Paragraph(text, title_style))
            elif line.startswith('## '):
                text = line.lstrip('#').strip()
                story.append(Paragraph(text, heading_style))
            elif line.startswith('---'):
                story.append(Spacer(1, 0.3*inch))
            else:
                # Regular paragraph
                # Escape special characters for ReportLab
                line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(line, body_style))
                
        # Build PDF
        doc.build(story)
        
        self.log_progress(f"Saved PDF: {output_path}")
        return output_path
        
    async def _export_docx(
        self,
        report: str,
        metadata: Dict,
        filename: str
    ) -> Path:
        """Export to DOCX using python-docx."""
        
        output_path = config.REPORTS_DIR / f"{filename}.docx"
        
        # Create document
        doc = Document()
        
        # Add title
        title = metadata.get("title", "Research Report")
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        date_para = doc.add_paragraph(f"Generated: {metadata.get('date', 'N/A')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Parse and add content
        for line in report.split('\n'):
            line = line.strip()
            
            if not line:
                continue
                
            if line.startswith('# '):
                doc.add_heading(line.lstrip('#').strip(), 1)
            elif line.startswith('## '):
                doc.add_heading(line.lstrip('#').strip(), 2)
            elif line.startswith('### '):
                doc.add_heading(line.lstrip('#').strip(), 3)
            elif line.startswith('---'):
                doc.add_paragraph()
            else:
                para = doc.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
        # Save
        doc.save(output_path)
        
        self.log_progress(f"Saved DOCX: {output_path}")
        return output_path
